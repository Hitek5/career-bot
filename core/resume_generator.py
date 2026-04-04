import json
import logging
import re
from pathlib import Path

from core.ai_engine import call_claude
from config import TEMPLATES_DIR, OUTPUT_DIR

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = """\
Ты — профессиональный составитель резюме. Создай адаптированное резюме под конкретную вакансию.

На основе профиля кандидата и анализа вакансии верни JSON:
{
    "full_name": "ИМЯ ФАМИЛИЯ",
    "target_position": "Целевая позиция (из вакансии)",
    "contacts": {"phone": "", "email": "", "telegram": "", "linkedin": ""},
    "summary": "2-3 предложения, продающие кандидата под ЭТУ вакансию",
    "experience": [
        {
            "company": "Компания",
            "position": "Должность",
            "period": "2020 — н.в.",
            "achievements": ["Достижение с цифрой"]
        }
    ],
    "education": [
        {"institution": "ВУЗ", "degree": "Степень", "field": "Направление", "year": "2008"}
    ],
    "skills": ["Навык1", "Навык2"],
    "certifications": ["Сертификат"],
    "languages": ["Русский — родной", "Английский — B2"]
}

Правила:
- Адаптируй порядок опыта и навыков под вакансию
- Акцентируй релевантные достижения
- Включай ключевые слова из вакансии
- Достижения ВСЕГДА с цифрами
- Максимум 4-5 мест работы, 3-4 достижения на место
- Summary заточено под вакансию
- Ответ ТОЛЬКО JSON"""


def _clean_json(text: str) -> str:
    import re
    m = re.search(r'```(?:json)?\s*\n(\{.*?\})\s*```', text, re.DOTALL)
    if m:
        return m.group(1).strip()
    m = re.search(r'(\{[\s\S]*\})', text)
    if m:
        return m.group(1).strip()
    return text.strip()


async def generate_resume_data(profile_json: str, analysis_json: str) -> dict:
    """Generate adapted resume content via AI."""
    # Extract enrichment data if available
    extra_context = ""
    try:
        profile_data = json.loads(profile_json)
        if profile_data.get("enriched_description"):
            extra_context += (
                f"\nОПИСАНИЕ КАНДИДАТА (из анализа предыдущих правок):\n"
                f"{profile_data['enriched_description']}\n"
            )
        if profile_data.get("preferred_style"):
            extra_context += (
                f"\nПРЕДПОЧТИТЕЛЬНЫЙ СТИЛЬ РЕЗЮМЕ:\n"
                f"{profile_data['preferred_style']}\n"
            )
    except (json.JSONDecodeError, TypeError):
        pass

    user_msg = (
        f"ПРОФИЛЬ КАНДИДАТА:\n{profile_json}\n\n"
        f"АНАЛИЗ ВАКАНСИИ:\n{analysis_json}\n"
    )
    if extra_context:
        user_msg += f"\nДОПОЛНИТЕЛЬНЫЙ КОНТЕКСТ:{extra_context}\n"
    user_msg += "\nСоздай адаптированное резюме под эту вакансию."

    response = await call_claude(SYSTEM_PROMPT, user_msg, max_tokens=4096)
    return json.loads(_clean_json(response))


def _build_contacts_html(contacts: dict) -> str:
    parts = []
    for key in ("phone", "email", "telegram", "linkedin"):
        v = contacts.get(key)
        if v:
            parts.append(v)
    return "<br>".join(parts)


def generate_pdf(resume_data: dict, output_name: str) -> Path:
    """Generate 1-page PDF resume using WeasyPrint."""
    from weasyprint import HTML

    template_path = TEMPLATES_DIR / "resume_1page.html"
    with open(template_path) as f:
        template = f.read()

    contacts = resume_data.get("contacts", {})

    # Experience HTML
    exp_html = ""
    for job in resume_data.get("experience", []):
        ach = "".join(f"<li>{a}</li>" for a in job.get("achievements", []))
        exp_html += (
            f'<div class="job">'
            f'<div class="job-header">'
            f'<div class="job-company">{job["company"]}</div>'
            f'<div class="job-period">{job["period"]}</div>'
            f'</div>'
            f'<div class="job-position">{job["position"]}</div>'
            f'<ul class="achievements">{ach}</ul>'
            f'</div>'
        )

    # Education HTML
    edu_html = ""
    for edu in resume_data.get("education", []):
        edu_html += (
            f'<div class="edu-item">'
            f'<span class="edu-inst">{edu["institution"]}</span> — '
            f'{edu.get("degree", "")} {edu.get("field", "")}'
            f'<span class="edu-year">{edu.get("year", "")}</span>'
            f'</div>'
        )

    skills_html = " &bull; ".join(resume_data.get("skills", []))

    certs = resume_data.get("certifications", [])
    certs_section = ""
    if certs:
        certs_section = (
            '<div class="section-title">Сертификаты</div>'
            f'<div class="info-line">{" &bull; ".join(certs)}</div>'
        )

    langs = resume_data.get("languages", [])
    langs_section = ""
    if langs:
        langs_section = (
            '<div class="section-title">Языки</div>'
            f'<div class="info-line">{" &bull; ".join(langs)}</div>'
        )

    html = (
        template
        .replace("{{FULL_NAME}}", resume_data.get("full_name", ""))
        .replace("{{TARGET_POSITION}}", resume_data.get("target_position", ""))
        .replace("{{CONTACTS}}", _build_contacts_html(contacts))
        .replace("{{SUMMARY}}", resume_data.get("summary", ""))
        .replace("{{EXPERIENCE}}", exp_html)
        .replace("{{EDUCATION}}", edu_html)
        .replace("{{SKILLS}}", skills_html)
        .replace("{{CERTIFICATIONS_SECTION}}", certs_section)
        .replace("{{LANGUAGES_SECTION}}", langs_section)
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{output_name}.pdf"
    HTML(string=html).write_pdf(str(output_path))
    return output_path


def generate_docx(resume_data: dict, output_name: str) -> Path:
    """Generate DOCX resume in HH-friendly format."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- Name ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(resume_data.get("full_name", ""))
    run.bold = True
    run.font.size = Pt(16)

    # --- Target position ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(resume_data.get("target_position", ""))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Contacts ---
    contacts = resume_data.get("contacts", {})
    parts = [contacts[k] for k in ("phone", "email", "telegram", "linkedin") if contacts.get(k)]
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(" | ".join(parts)).font.size = Pt(9)

    # --- Summary ---
    doc.add_paragraph()
    doc.add_paragraph(resume_data.get("summary", ""))

    # --- Experience ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("ОПЫТ РАБОТЫ")
    run.bold = True
    run.font.size = Pt(12)

    for job in resume_data.get("experience", []):
        p = doc.add_paragraph()
        p.add_run(job["company"]).bold = True
        p.add_run(f"  |  {job['period']}")

        p = doc.add_paragraph()
        p.add_run(job["position"]).italic = True

        for ach in job.get("achievements", []):
            doc.add_paragraph(ach, style="List Bullet")

    # --- Education ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("ОБРАЗОВАНИЕ")
    run.bold = True
    run.font.size = Pt(12)

    for edu in resume_data.get("education", []):
        p = doc.add_paragraph()
        p.add_run(edu["institution"]).bold = True
        p.add_run(f" — {edu.get('degree', '')} {edu.get('field', '')}, {edu.get('year', '')}")

    # --- Skills ---
    skills = resume_data.get("skills", [])
    if skills:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("НАВЫКИ").bold = True
        doc.add_paragraph(" • ".join(skills))

    # --- Certifications ---
    certs = resume_data.get("certifications", [])
    if certs:
        p = doc.add_paragraph()
        p.add_run("СЕРТИФИКАТЫ").bold = True
        doc.add_paragraph(" • ".join(certs))

    # --- Languages ---
    langs = resume_data.get("languages", [])
    if langs:
        p = doc.add_paragraph()
        p.add_run("ЯЗЫКИ").bold = True
        doc.add_paragraph(" • ".join(langs))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{output_name}.docx"
    doc.save(str(output_path))
    return output_path
