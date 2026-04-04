"""Universal document parser — converts any office format to text.

Supported: .pdf, .docx, .doc, .odt, .rtf, .txt
Uses LibreOffice for conversion of non-native formats.
"""

import subprocess
import tempfile
from pathlib import Path

from parsers.pdf_parser import parse_pdf
from parsers.docx_parser import parse_docx


# Formats that LibreOffice can convert to docx
LIBREOFFICE_FORMATS = {".doc", ".odt", ".rtf"}
NATIVE_FORMATS = {".pdf", ".docx", ".txt"}
ALL_SUPPORTED = NATIVE_FORMATS | LIBREOFFICE_FORMATS


def _convert_to_docx(file_path: Path) -> Path:
    """Convert any office format to .docx via LibreOffice."""
    with tempfile.TemporaryDirectory() as tmp_dir:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--norestore",
                "--convert-to",
                "docx",
                str(file_path),
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            timeout=120,
        )
        stderr = result.stderr.decode(errors="replace")
        if result.returncode != 0 or "Error:" in stderr:
            raise RuntimeError(f"LibreOffice conversion failed: {stderr}")

        converted = Path(tmp_dir) / (file_path.stem + ".docx")
        if not converted.exists():
            # LibreOffice sometimes names output differently
            docx_files = list(Path(tmp_dir).glob("*.docx"))
            if docx_files:
                converted = docx_files[0]
            else:
                raise FileNotFoundError(
                    f"No .docx output after LibreOffice conversion of {file_path.name}"
                )

        # Move out of tmp dir
        final = file_path.with_suffix(".converted.docx")
        converted.rename(final)
        return final


def _convert_to_txt(file_path: Path) -> Path:
    """Convert any office format to .txt via LibreOffice (fallback)."""
    with tempfile.TemporaryDirectory() as tmp_dir:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--norestore",
                "--convert-to",
                "txt:Text",
                str(file_path),
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError("LibreOffice txt conversion failed")

        txt_files = list(Path(tmp_dir).glob("*.txt"))
        if not txt_files:
            raise FileNotFoundError("No .txt output after conversion")

        final = file_path.with_suffix(".converted.txt")
        txt_files[0].rename(final)
        return final


def parse_document(file_path: str | Path) -> str:
    """Parse any supported document format and return extracted text.

    Raises RuntimeError/FileNotFoundError if conversion fails.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    cleanup_files: list[Path] = []

    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()

        if ext == ".pdf":
            return parse_pdf(str(file_path))

        if ext == ".docx":
            return parse_docx(str(file_path))

        if ext in LIBREOFFICE_FORMATS:
            # For .doc files, try antiword/catdoc first (faster, no LO lock issues)
            if ext == ".doc":
                for tool in ("antiword", "catdoc"):
                    try:
                        r = subprocess.run(
                            [tool, str(file_path)],
                            capture_output=True, timeout=30,
                        )
                        if r.returncode == 0:
                            text = r.stdout.decode("utf-8", errors="replace").strip()
                            if text:
                                return text
                    except (FileNotFoundError, subprocess.TimeoutExpired):
                        continue

            # Convert to docx first, then parse with python-docx
            try:
                docx_path = _convert_to_docx(file_path)
                cleanup_files.append(docx_path)
                from docx import Document

                doc = Document(str(docx_path))
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            c.text.strip() for c in row.cells if c.text.strip()
                        )
                        if row_text:
                            parts.append(row_text)
                content = "\n".join(parts)
                if content.strip():
                    return content
            except Exception:
                pass  # Fall through to txt fallback

            # Fallback: convert to txt directly
            try:
                txt_path = _convert_to_txt(file_path)
                cleanup_files.append(txt_path)
                with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
            except Exception as e:
                raise RuntimeError(
                    f"Не удалось обработать {file_path.name} ({ext}). "
                    f"Попробуйте сохранить как .docx или .pdf и загрузить снова."
                )

        raise ValueError(f"Неподдерживаемый формат: {ext}")

    finally:
        for f in cleanup_files:
            if f.exists():
                f.unlink(missing_ok=True)
