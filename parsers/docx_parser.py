import subprocess
import tempfile
from pathlib import Path

from docx import Document


def _convert_doc_to_docx(doc_path: str | Path) -> Path:
    """Convert .doc (Word 97-2003) to .docx via LibreOffice."""
    doc_path = Path(doc_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                str(doc_path),
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr.decode(errors='replace')}"
            )
        converted = Path(tmp_dir) / (doc_path.stem + ".docx")
        if not converted.exists():
            raise FileNotFoundError(
                f"Converted file not found after LibreOffice conversion: {converted}"
            )
        # Move to same dir as original so it survives tmp cleanup
        final = doc_path.with_suffix(".docx")
        converted.rename(final)
        return final


def parse_docx(file_path: str | Path) -> str:
    """Extract text from DOCX/DOC file (paragraphs + tables).

    If the file is .doc (old binary format), it's converted to .docx first
    via LibreOffice.
    """
    file_path = Path(file_path)
    cleanup_converted = False

    if file_path.suffix.lower() == ".doc":
        file_path = _convert_doc_to_docx(file_path)
        cleanup_converted = True

    try:
        doc = Document(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    c.text.strip() for c in row.cells if c.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    finally:
        if cleanup_converted and file_path.exists():
            file_path.unlink(missing_ok=True)
